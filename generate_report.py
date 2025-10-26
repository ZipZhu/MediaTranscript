"""将转录文本与摘要内容合并生成报告（DOCX 或 PDF）。

示例：
    python generate_report.py \
        --transcript transcript.txt \
        --summary summary.txt \
        --output report.docx

    python generate_report.py \
        --transcript transcript.txt \
        --summary summary.txt \
        --output report.pdf --format pdf
"""

from __future__ import annotations

import argparse
from datetime import datetime
from pathlib import Path
from typing import Literal

from docx import Document
from docx.enum.text import WD_BREAK
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


ReportFormat = Literal["docx", "pdf"]


def read_text_file(path: Path) -> str:
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {path}")
    return path.read_text(encoding="utf-8").strip()


def normalize_output_path(output: Path, report_format: ReportFormat) -> Path:
    suffix = output.suffix.lower()
    expected = f".{report_format}"
    if suffix != expected:
        output = output.with_suffix(expected)
    output.parent.mkdir(parents=True, exist_ok=True)
    return output


def build_report_title() -> str:
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
    return f"MediaTranscript 摘要报告 ({timestamp})"


def generate_docx(
    transcript: str,
    summary: str,
    output_path: Path,
) -> None:
    document = Document()

    document.add_heading(build_report_title(), level=1)

    document.add_heading("摘要", level=2)
    document.add_paragraph(summary)

    document.add_page_break()

    document.add_heading("全文转录", level=2)

    # Whisper 生成的文本通常较长，按段落划分显示更易阅读。
    for paragraph in transcript.splitlines():
        if paragraph.strip():
            document.add_paragraph(paragraph.strip())
        else:
            document.add_paragraph().add_run().add_break(WD_BREAK.LINE)

    document.save(output_path)


def generate_pdf(
    transcript: str,
    summary: str,
    output_path: Path,
) -> None:
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        name="TitleStyle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=18,
        textColor=colors.HexColor("#2C3E50"),
        spaceAfter=18,
    )

    section_style = ParagraphStyle(
        name="SectionHeading",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=14,
        textColor=colors.HexColor("#1F618D"),
        spaceAfter=12,
    )

    body_style = ParagraphStyle(
        name="BodyText",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=11,
        leading=16,
        spaceAfter=8,
    )

    story = []
    story.append(Paragraph(build_report_title(), title_style))
    story.append(Paragraph("摘要", section_style))

    for para in summary.splitlines():
        story.append(Paragraph(para or "\u00a0", body_style))
    story.append(Spacer(1, 1 * cm))

    story.append(Paragraph("全文转录", section_style))
    for para in transcript.splitlines():
        story.append(Paragraph(para or "\u00a0", body_style))

    doc = SimpleDocTemplate(str(output_path), pagesize=A4, leftMargin=2 * cm, rightMargin=2 * cm)
    doc.build(story)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="合并转录与摘要，生成 DOCX 或 PDF 报告。")
    parser.add_argument("--transcript", required=True, help="转录文本文件路径")
    parser.add_argument("--summary", required=True, help="摘要文本文件路径")
    parser.add_argument(
        "--output",
        required=True,
        help="输出报告路径（后缀可省略，将按 --format 自动补全）",
    )
    parser.add_argument(
        "--format",
        choices=["docx", "pdf"],
        default="docx",
        help="报告文件格式，默认 docx",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()

    transcript_path = Path(args.transcript).expanduser().resolve()
    summary_path = Path(args.summary).expanduser().resolve()
    output_path = Path(args.output).expanduser().resolve()
    report_format: ReportFormat = args.format

    transcript = read_text_file(transcript_path)
    summary = read_text_file(summary_path)

    output_path = normalize_output_path(output_path, report_format)

    if report_format == "docx":
        generate_docx(transcript, summary, output_path)
    else:
        generate_pdf(transcript, summary, output_path)

    print(f"报告已生成: {output_path}")


if __name__ == "__main__":
    main()

